# File: analyzer.py
import os
import re
from typing import List, Dict, Any

from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document
from azure.ai.textanalytics import TextAnalyticsClient
from azure.core.credentials import AzureKeyCredential

load_dotenv()

# ============================================================
#  Azure Client Setup
# ============================================================

def _get_client() -> TextAnalyticsClient:
    endpoint = os.getenv("AZURE_LANGUAGE_ENDPOINT")
    key = os.getenv("AZURE_LANGUAGE_KEY")

    if not endpoint or not key:
        raise ValueError("Azure credentials missing in .env")

    return TextAnalyticsClient(endpoint, AzureKeyCredential(key))


# ============================================================
#  File Text Extraction (PDF + DOCX + tables)
# ============================================================

def extract_text(file_path: str) -> str:
    """
    Extracts textual content from PDF or DOCX.
    For DOCX, also reads table cells (useful for modern resume templates).
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    _, ext = os.path.splitext(file_path)
    text = ""

    if ext.lower() == ".pdf":
        reader = PdfReader(file_path)
        for page in reader.pages:
            raw = page.extract_text() or ""
            # Normalize spacing a bit
            raw = raw.replace("\t", " ")
            while "  " in raw:
                raw = raw.replace("  ", " ")
            text += raw + "\n"

    elif ext.lower() == ".docx":
        doc = Document(file_path)
        parts = [p.text for p in doc.paragraphs]

        # Include table text too
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)

        text = "\n".join(parts)

    else:
        raise ValueError("Unsupported file type. Only PDF and DOCX are allowed.")

    return text


# ============================================================
#  Helpers: Email, Sections, Skills, Experience
# ============================================================

def try_recover_email_from_spaced_text(text: str) -> str | None:
    """
    Tries to recover emails where characters are spaced out like:
    y e s h w a n t h 3 0 0 6 @ g m a i l . c o m
    """
    compact = text.replace(" ", "")
    m = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", compact)
    return m.group(0) if m else None


def detect_missing_sections(resume_text: str) -> List[str]:
    """
    Robust detection for presence of:
    Skills, Education, Experience, Projects, Certification.

    Handles cases where headings come out as 'S K I L L S' in PDFs.
    """
    text = resume_text.lower()
    compact = text.replace(" ", "")

    section_keywords = {
        "Skills": [
            "skills", "technical skills", "skill set", "skills & abilities", "technologies"
        ],
        "Education": [
            "education", "academic background", "educational qualification",
            "academics", "qualification", "b.tech", "bachelor", "degree"
        ],
        "Experience": [
            "experience", "work experience", "professional experience",
            "employment history", "career history", "internship", "intern", "work history"
        ],
        "Projects": [
            "projects", "project work", "academic projects", "personal projects",
            "project experience"
        ],
        "Certification": [
            "certifications", "certification", "courses", "training"
        ],
    }

    missing = []

    for section_name, keys in section_keywords.items():
        found = False
        for k in keys:
            k_low = k.lower()
            # check in normal text OR in space-removed text
            if k_low in text or k_low.replace(" ", "") in compact:
                found = True
                break
        if not found:
            missing.append(section_name)

    return missing


COMMON_SKILLS = [
    "python", "java", "javascript", "typescript", "c", "c++", "c#", "go", "rust",
    "html", "css", "react", "react.js", "angular", "vue", "node", "node.js",
    "express", "express.js", "django", "flask",
    "sql", "mysql", "postgresql", "mongodb",
    "azure", "aws", "gcp", "docker", "kubernetes", "git", "github",
    "pandas", "numpy", "tensorflow", "pytorch",
]


def detect_skills(text: str, key_phrases: List[str]) -> List[str]:
    found = set()
    low = text.lower()
    compact = low.replace(" ", "")

    # Known skills from text
    for s in COMMON_SKILLS:
        if s in low or s.replace(" ", "") in compact:
            found.add(s.upper() if len(s) <= 3 else s.title())

    # From key phrases if any match directly
    for kp in key_phrases:
        kpl = kp.lower()
        if kpl in COMMON_SKILLS:
            found.add(kp.upper() if len(kpl) <= 3 else kp.title())

    return sorted(found)


def extract_experience_lines(text: str) -> List[str]:
    """
    Extracts lines that likely describe experience
    using years and role-related keywords.
    """
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    exp_lines = []

    keywords = [
        "experience", "work experience", "professional experience",
        "intern", "internship", "developer", "engineer",
        "company", "organisation", "organization", "trainee"
    ]

    for line in lines:
        low = line.lower()
        compact = low.replace(" ", "")

        year_found = any(str(y) in compact for y in range(2015, 2031))
        kw_found = any(k in low or k.replace(" ", "") in compact for k in keywords)

        if year_found or kw_found:
            exp_lines.append(line)

    return exp_lines[:10]


# ============================================================
#  Core Resume Analysis (email/phone/keyphrases/entities)
#   (Name is simple: Azure only, otherwise None)
# ============================================================

def analyze_resume_text(text: str) -> Dict[str, Any]:
    client = _get_client()
    docs = [text]

    # Key phrases
    try:
        kp_resp = client.extract_key_phrases(docs)[0]
        key_phrases = kp_resp.key_phrases if not kp_resp.is_error else []
    except Exception as e:
        print("Key phrase error:", e)
        key_phrases = []

    # Entities
    entities: List[Dict[str, Any]] = []
    try:
        ent_resp = client.recognize_entities(docs)[0]
        if not ent_resp.is_error:
            for e in ent_resp.entities:
                entities.append(
                    {
                        "text": e.text,
                        "category": e.category,
                        "sub": e.subcategory,
                        "score": e.confidence_score,
                    }
                )
    except Exception as e:
        print("Entity error:", e)

    # Name (only from Azure Person entity; no extra heuristics)
    name = next(
        (e["text"] for e in entities if e.get("category") == "Person"),
        None,
    )

    # Email
    email_match = re.search(
        r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text
    )
    if email_match:
        email = email_match.group(0)
    else:
        email = try_recover_email_from_spaced_text(text)

    # Phone
    phone_match = re.search(
        r"\+?\d[\d\s\-]{8,}", text
    )
    if phone_match:
        raw_phone = phone_match.group(0)
        # remove everything except digits and leading +
        phone = re.sub(r"[^\d+]", "", raw_phone)
    else:
        phone = None

    return {
        "raw_text": text,
        "key_phrases": key_phrases,
        "entities": entities,
        "name": name,      # may be None → "Not detected" in UI
        "email": email,
        "phone": phone,
    }


# ============================================================
#  JD Keyword Extraction + Matching
# ============================================================

def extract_jd_keywords(jd: str) -> List[str]:
    client = _get_client()
    try:
        resp = client.extract_key_phrases([jd])[0]
        if resp.is_error:
            return []
        return [k.strip().lower() for k in resp.key_phrases if k.strip()]
    except Exception as e:
        print("JD key phrase error:", e)
        return []


def analyze_resume_with_job(file_path: str, job_description: str) -> Dict[str, Any]:
    """
    Main high-level function used by Flask:
      - Extracts text
      - Analyzes resume (name, email, phone, key phrases, entities)
      - Detects skills
      - Extracts experience lines
      - Checks missing sections (Skills, Education, Experience, Projects, Certification)
      - Extracts JD keywords, computes match %, missing/matched keywords
      - Produces suggestions
    """
    text = extract_text(file_path)
    if not text.strip():
        raise ValueError("The resume appears to be empty or unreadable.")

    base = analyze_resume_text(text)

    # Skills & Experience
    skills_detected = detect_skills(text, base["key_phrases"])
    experience_summary = extract_experience_lines(text)

    # Section presence (robust, spacing-safe)
    missing_sections = detect_missing_sections(text)

    # JD matching
    jd_keywords = extract_jd_keywords(job_description)
    resume_lower_compact = text.lower().replace(" ", "")  # compress spaces for matching

    matched_keywords = [k for k in jd_keywords if k.replace(" ", "") in resume_lower_compact]
    missing_keywords = [k for k in jd_keywords if k.replace(" ", "") not in resume_lower_compact]

    match_percentage = int((len(matched_keywords) / len(jd_keywords)) * 100) if jd_keywords else 0

    # Suggestions
    suggestions: List[str] = []

    if missing_keywords:
        suggestions.append(
            "Consider including or emphasizing these job-related keywords if relevant: "
            + ", ".join(sorted(set(missing_keywords[:10])))
        )

    if missing_sections:
        suggestions.append(
            "Your resume may be missing or not clearly highlighting these sections: "
            + ", ".join(missing_sections)
        )

    if match_percentage < 60:
        suggestions.append(
            "Your overall match score is below 60%. Tailor your resume more closely to the job description."
        )
    elif match_percentage < 80:
        suggestions.append(
            "Your match score is decent. You can still improve it by emphasizing key skills and achievements from the JD."
        )
    else:
        suggestions.append(
            "Great! Your resume is strongly aligned with this job description."
        )

    if not suggestions:
        suggestions.append("Your resume looks solid. Focus on measurable achievements and clarity.")

    # Final combined result – matches what your HTML expects
    result: Dict[str, Any] = {
        "name": base["name"],
        "email": base["email"],
        "phone": base["phone"],
        "skills_detected": skills_detected,
        "experience_summary": experience_summary,
        "match_percentage": match_percentage,
        "matched_keywords": matched_keywords,
        "missing_keywords": missing_keywords,
        "missing_sections": missing_sections,
        "suggestions": suggestions,
        "job_description": job_description,
        # extras if needed
        "key_phrases": base["key_phrases"],
        "entities": base["entities"],
        "raw_text": base["raw_text"],
    }

    return result
